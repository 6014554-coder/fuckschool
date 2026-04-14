"""
ГОСТ-форматирование .docx документов.

Применяемые правила по умолчанию (ГОСТ 7.32-2017):
- Шрифт: Times New Roman, 14 pt
- Межстрочный интервал: 1.5
- Поля: левое 30 мм, правое 10 мм, верхнее 20 мм, нижнее 20 мм
- Отступ первой строки абзаца: 1.25 см
- Выравнивание основного текста: по ширине
- Заголовки 1 уровня: 14 pt, жирный, КАПСЛОК, по центру, без отступа, с новой страницы
- Заголовки 2 уровня: 14 pt, жирный, по левому краю, без отступа
- Заголовки 3 уровня: 14 pt, жирный, по левому краю, без отступа
- Нумерация страниц: снизу по центру, без точки

Если передан config (из analyzer.analyze_example), используются параметры из него.
"""

import re
from io import BytesIO

from docx import Document
from docx.shared import Pt, Mm, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# ---------------------------------------------------------------------------
# Константы ГОСТ (defaults)
# ---------------------------------------------------------------------------

FONT_NAME = "Times New Roman"
FONT_SIZE = Pt(14)

MARGIN_LEFT   = Mm(30)
MARGIN_RIGHT  = Mm(10)
MARGIN_TOP    = Mm(20)
MARGIN_BOTTOM = Mm(20)

FIRST_LINE_INDENT = Cm(1.25)

HEADING_CONFIG = {
    1: {"bold": True,  "italic": False, "caps": True,  "align": WD_ALIGN_PARAGRAPH.CENTER,
        "space_before": Pt(0),  "space_after": Pt(12), "page_break": True,
        "first_line_indent": Pt(0)},
    2: {"bold": True,  "italic": False, "caps": False, "align": WD_ALIGN_PARAGRAPH.LEFT,
        "space_before": Pt(12), "space_after": Pt(6),  "page_break": False,
        "first_line_indent": Pt(0)},
    3: {"bold": True,  "italic": False, "caps": False, "align": WD_ALIGN_PARAGRAPH.LEFT,
        "space_before": Pt(8),  "space_after": Pt(4),  "page_break": False,
        "first_line_indent": Pt(0)},
}

# Ключевые слова для заголовков 1 уровня без нумерации
UNNUMBERED_HEADINGS = {
    "ВВЕДЕНИЕ", "INTRODUCTION",
    "ЗАКЛЮЧЕНИЕ", "CONCLUSION",
    "ВЫВОДЫ", "SUMMARY",
    "АННОТАЦИЯ", "ABSTRACT",
    "СОДЕРЖАНИЕ", "ОГЛАВЛЕНИЕ",
    "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ",
    "СПИСОК ЛИТЕРАТУРЫ",
    "СПИСОК ИСПОЛЬЗОВАННОЙ ЛИТЕРАТУРЫ",
    "REFERENCES",
    "ПРИЛОЖЕНИЕ", "ПРИЛОЖЕНИЯ",
    "БЛАГОДАРНОСТИ",
}


# ---------------------------------------------------------------------------
# Поиск начала содержательной части (после титульника)
# ---------------------------------------------------------------------------

# Маркеры, с которых начинается содержательная часть
_CONTENT_START_WORDS = {
    "СОДЕРЖАНИЕ", "ОГЛАВЛЕНИЕ",
    "АННОТАЦИЯ", "ABSTRACT",
    "ВВЕДЕНИЕ", "INTRODUCTION",
}

def _find_content_start(paragraphs) -> int:
    """
    Возвращает индекс первого параграфа содержательной части.
    Ищет СОДЕРЖАНИЕ / АННОТАЦИЯ / ВВЕДЕНИЕ / нумерованный заголовок "1.".
    Если не нашли — возвращает 0 (форматируем всё).
    """
    for i, para in enumerate(paragraphs):
        text = para.text.strip()
        if not text:
            continue
        upper = text.upper().rstrip(" .")
        if upper in _CONTENT_START_WORDS:
            return i
        # Нумерованный раздел: "1 Название" или "1. Название"
        if re.match(r'^1\.?\s+\S', text):
            return i
    return 0


# ---------------------------------------------------------------------------
# Определение типа параграфа
# ---------------------------------------------------------------------------

def _get_heading_level_by_style(para):
    """Заголовок по стилю Word (Heading 1/2/3)."""
    style_name = para.style.name if para.style else ""
    for level in (1, 2, 3):
        if style_name in (f"Heading {level}", f"Заголовок {level}"):
            return level
    return None


def _get_heading_level_by_content(para):
    """
    Определяет уровень заголовка по содержимому параграфа.
    Работает когда все параграфы написаны стилем Normal/None.
    """
    text = para.text.strip()
    if not text or len(text) < 2:
        return None

    is_bold = any(run.font.bold for run in para.runs)
    if not is_bold:
        return None

    if re.match(r'^\d+\.\d+\.\d+', text):
        return 3
    if re.match(r'^\d+\.\d+', text):
        return 2
    if re.match(r'^\d+\.?\s+\S', text):
        return 1

    text_upper = text.upper()
    if text_upper in UNNUMBERED_HEADINGS:
        return 1
    if len(text) >= 4 and text == text_upper and not text[0].isdigit():
        return 1

    return None


def _is_body_text(para) -> bool:
    style_name = para.style.name if para.style else ""
    body_styles = {"Normal", "Body Text", "Основной текст", ""}
    return style_name in body_styles or style_name.startswith("Normal")


def _is_table_caption(para) -> bool:
    text = para.text.strip().upper()
    return text.startswith("ТАБЛИЦА") or text.startswith("ТАБЛИЦЯ")


def _is_figure_caption(para) -> bool:
    text = para.text.strip()
    return (text.lower().startswith("рис.") or
            text.lower().startswith("рисунок") or
            text.lower().startswith("fig."))


def _is_list_item(para) -> bool:
    style_name = para.style.name if para.style else ""
    return "List" in style_name or "Список" in style_name


# ---------------------------------------------------------------------------
# Применение форматирования
# ---------------------------------------------------------------------------

def _set_font(run, font_name, font_size, bold=None, italic=None, caps=None):
    run.font.name = font_name
    run.font.size = font_size
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    if caps is not None:
        run.font.all_caps = caps
    run.font.color.rgb = RGBColor(0, 0, 0)


def _apply_heading(para, level: int, cfg: dict, font_name: str, font_size):
    h = cfg["headings"].get(level)
    if not h:
        return

    fmt = para.paragraph_format
    fmt.alignment         = h["alignment"]
    fmt.first_line_indent = Cm(h["first_line_indent_cm"])
    fmt.space_before      = Pt(h["space_before_pt"])
    fmt.space_after       = Pt(h["space_after_pt"])
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    fmt.page_break_before = h["page_break"]

    for run in para.runs:
        _set_font(run, font_name, font_size,
                  bold=h["bold"], italic=h["italic"], caps=h["caps"])


def _apply_body(para, cfg: dict, font_name: str, font_size):
    b = cfg["body"]
    fmt = para.paragraph_format
    fmt.alignment         = b["alignment"]
    fmt.first_line_indent = Cm(b["first_line_indent_cm"])
    fmt.space_before      = Pt(b["space_before_pt"])
    fmt.space_after       = Pt(b["space_after_pt"])
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    fmt.page_break_before = False

    for run in para.runs:
        _set_font(run, font_name, font_size, bold=False, italic=False, caps=False)


def _apply_table_caption(para, font_name, font_size):
    fmt = para.paragraph_format
    fmt.alignment         = WD_ALIGN_PARAGRAPH.LEFT
    fmt.first_line_indent = Pt(0)
    fmt.space_before      = Pt(12)
    fmt.space_after       = Pt(3)
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    for run in para.runs:
        _set_font(run, font_name, font_size, bold=False, italic=False, caps=False)


def _apply_figure_caption(para, font_name, font_size):
    fmt = para.paragraph_format
    fmt.alignment         = WD_ALIGN_PARAGRAPH.CENTER
    fmt.first_line_indent = Pt(0)
    fmt.space_before      = Pt(3)
    fmt.space_after       = Pt(12)
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    for run in para.runs:
        _set_font(run, font_name, font_size, bold=False, italic=False, caps=False)


def _apply_list_item(para, font_name, font_size):
    fmt = para.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    fmt.first_line_indent = Pt(0)
    fmt.left_indent       = Cm(1.25)
    fmt.space_before      = Pt(0)
    fmt.space_after       = Pt(0)
    for run in para.runs:
        _set_font(run, font_name, font_size, bold=False, italic=False, caps=False)


# ---------------------------------------------------------------------------
# Нумерация страниц
# ---------------------------------------------------------------------------

def _add_page_numbers(doc: Document, font_name: str, font_size):
    """Добавляет нумерацию страниц снизу по центру."""
    section = doc.sections[0]
    footer = section.footer
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.clear()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = para.add_run()
    run.font.name = font_name
    run.font.size = font_size

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    instr_run = para.add_run()
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = " PAGE "
    instr_run._r.append(instrText)

    fld_end_run = para.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    fld_end_run._r.append(fld_end)


# ---------------------------------------------------------------------------
# Построение конфига из ГОСТ-defaults
# ---------------------------------------------------------------------------

def _default_config() -> dict:
    return {
        "font_name":    FONT_NAME,
        "font_size_pt": 14.0,
        "margins": {
            "left_mm":   30.0,
            "right_mm":  10.0,
            "top_mm":    20.0,
            "bottom_mm": 20.0,
        },
        "body": {
            "first_line_indent_cm": 1.25,
            "alignment":            WD_ALIGN_PARAGRAPH.JUSTIFY,
            "space_before_pt":      0.0,
            "space_after_pt":       0.0,
            "line_spacing":         WD_LINE_SPACING.ONE_POINT_FIVE,
        },
        "headings": {
            1: {"bold": True,  "italic": False, "caps": True,
                "alignment": WD_ALIGN_PARAGRAPH.CENTER,
                "first_line_indent_cm": 0.0,
                "space_before_pt": 0.0, "space_after_pt": 12.0,
                "page_break": True},
            2: {"bold": True,  "italic": False, "caps": False,
                "alignment": WD_ALIGN_PARAGRAPH.LEFT,
                "first_line_indent_cm": 0.0,
                "space_before_pt": 12.0, "space_after_pt": 6.0,
                "page_break": False},
            3: {"bold": True,  "italic": False, "caps": False,
                "alignment": WD_ALIGN_PARAGRAPH.LEFT,
                "first_line_indent_cm": 0.0,
                "space_before_pt": 8.0, "space_after_pt": 4.0,
                "page_break": False},
        },
    }


# ---------------------------------------------------------------------------
# Основная функция
# ---------------------------------------------------------------------------

def format_document(docx_bytes: bytes, config: dict = None) -> bytes:
    """
    Принимает байты .docx, возвращает байты отформатированного .docx.

    config — опциональный словарь из analyzer.analyze_example().
             Если None, применяются стандартные правила ГОСТ.
    """
    if config is None:
        config = _default_config()

    # Заполняем недостающие уровни заголовков из defaults
    default = _default_config()
    for lvl in (1, 2, 3):
        if lvl not in config.get("headings", {}):
            config.setdefault("headings", {})[lvl] = default["headings"][lvl]

    font_name = config.get("font_name", FONT_NAME)
    font_size = Pt(config.get("font_size_pt", 14.0))
    margins   = config.get("margins", default["margins"])

    doc = Document(BytesIO(docx_bytes))

    # 1. Поля страницы
    for section in doc.sections:
        section.left_margin   = Mm(margins["left_mm"])
        section.right_margin  = Mm(margins["right_mm"])
        section.top_margin    = Mm(margins["top_mm"])
        section.bottom_margin = Mm(margins["bottom_mm"])

    # 2. Находим начало содержательной части (пропускаем титульник)
    all_paras = doc.paragraphs
    content_start = _find_content_start(all_paras)

    # 3. Обход параграфов: до content_start — не трогаем, после — форматируем
    for idx, para in enumerate(all_paras):
        if idx < content_start:
            continue
        if not para.text.strip():
            continue

        heading_level = (_get_heading_level_by_style(para) or
                         _get_heading_level_by_content(para))

        if heading_level is not None:
            _apply_heading(para, heading_level, config, font_name, font_size)

        elif _is_list_item(para):
            _apply_list_item(para, font_name, font_size)

        elif _is_table_caption(para):
            _apply_table_caption(para, font_name, font_size)

        elif _is_figure_caption(para):
            _apply_figure_caption(para, font_name, font_size)

        elif _is_body_text(para):
            _apply_body(para, config, font_name, font_size)

        else:
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            for run in para.runs:
                run.font.name = font_name
                run.font.size = font_size

    # 3. Нумерация страниц
    _add_page_numbers(doc, font_name, font_size)

    # 4. Сохраняем в байты
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
